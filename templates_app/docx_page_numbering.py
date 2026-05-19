"""
Injeta "Página X de Y" no rodapé de cada seção de um .docx.

Usa o formato OOXML canônico de complex field (fldChar begin / instrText /
fldChar separate / cache / fldChar end) para PAGE e NUMPAGES. Esse é o
único formato que LibreOffice headless avalia de forma consistente ao
converter para PDF — `<w:fldSimple>` com run-cache e IF aninhado deixam
o rodapé vazio em versões 7.x.

Ambas funções (`add_page_numbering_simple` e `add_page_numbering_conditional`)
hoje produzem o mesmo resultado — "Página X de Y" sempre visível. A versão
"conditional" foi mantida apenas por retrocompatibilidade de import; o IF
aninhado original quebrava no caminho PDF e em viewers fora do Word.

Trade-off conhecido: docs de 1 página exibem "Página 1 de 1". Onde isso
não é desejado, o caller deve detectar a contagem antes (ver
`_convert_with_conditional_page_numbering` em views.py).

Em caso de erro, retornam os bytes originais sem quebrar o pipeline.
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _footer_has_page_field(footer) -> bool:
    """True se o rodapé já tem o padrão 'X de Y' (campo NUMPAGES presente)."""
    xml = footer._element.xml if hasattr(footer._element, "xml") else ""
    if not xml:
        for p in footer.paragraphs:
            text = getattr(p._p, "xml", "")
            if "NUMPAGES" in text:
                return True
        return False
    return "NUMPAGES" in xml


def _make_run_with_child(child):
    r = OxmlElement("w:r")
    r.append(child)
    return r


def _fld_char(char_type: str):
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), char_type)
    return _make_run_with_child(fc)


def _instr_text(text: str):
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = text
    return _make_run_with_child(it)


def _text_run(text: str):
    t = OxmlElement("w:t")
    t.text = text
    if not text or text != text.strip():
        t.set(qn("xml:space"), "preserve")
    return _make_run_with_child(t)


def _append_complex_field(p_el, instr: str, cache_text: str = "1") -> None:
    """Anexa um campo complex (fldChar begin / instrText / separate / cache / end).

    Forma canônica que LibreOffice headless avalia corretamente ao converter
    para PDF. `cache_text` é o valor exibido caso o renderer não atualize os
    campos — '1' é seguro pra PAGE e NUMPAGES.
    """
    p_el.append(_fld_char("begin"))
    p_el.append(_instr_text(instr))
    p_el.append(_fld_char("separate"))
    p_el.append(_text_run(cache_text))
    p_el.append(_fld_char("end"))


def _append_simple_paragraph(footer):
    """'Página { PAGE } de { NUMPAGES }' alinhado à direita, sempre visível."""
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_el = p._p
    p_el.append(_text_run("Página "))
    _append_complex_field(p_el, "PAGE")
    p_el.append(_text_run(" de "))
    _append_complex_field(p_el, "NUMPAGES")


# Alias mantido por retrocompatibilidade. O IF aninhado anterior quebrava em
# qualquer viewer fora do Word — agora ambas as estratégias produzem o mesmo
# parágrafo simple com complex fields.
_append_conditional_paragraph = _append_simple_paragraph


def _section_has_different_first_page(section) -> bool:
    sect_pr = section._sectPr
    return sect_pr.find(qn("w:titlePg")) is not None


def _add(docx_bytes: bytes, append_fn) -> bytes:
    doc = Document(BytesIO(docx_bytes))

    for section in doc.sections:
        footer = section.footer
        if not _footer_has_page_field(footer):
            append_fn(footer)

        if _section_has_different_first_page(section):
            first_footer = section.first_page_footer
            if not _footer_has_page_field(first_footer):
                append_fn(first_footer)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def add_page_numbering_simple(docx_bytes: bytes) -> bytes:
    """Injeta 'Página X de Y' sempre visível. Use após detectar multi-página."""
    try:
        return _add(docx_bytes, _append_simple_paragraph)
    except Exception:
        return docx_bytes


def add_page_numbering_conditional(docx_bytes: bytes) -> bytes:
    """Mantida por retrocompat — hoje é equivalente a add_page_numbering_simple.
    Sempre injeta 'Página X de Y'. Docs de 1 página exibem 'Página 1 de 1'."""
    try:
        return _add(docx_bytes, _append_conditional_paragraph)
    except Exception:
        return docx_bytes


# Alias mantido por retrocompat com chamadas existentes.
add_page_numbering = add_page_numbering_simple
