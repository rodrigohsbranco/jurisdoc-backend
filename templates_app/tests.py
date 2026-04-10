from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from django.test import SimpleTestCase
from docx import Document
from docxtpl import DocxTemplate

from common.jinja_env import build_env
from templates_app.docx_jinja_normalizer import normalize_docx_jinja_runs


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


class DocxJinjaNormalizerTests(SimpleTestCase):
    def test_keeps_bold_when_placeholder_is_split_across_runs(self):
        with TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "template.docx"

            doc = Document()
            p = doc.add_paragraph()
            p.add_run("{{")
            styled = p.add_run("nome")
            styled.bold = True
            p.add_run("}}")
            doc.save(src)

            normalized = normalize_docx_jinja_runs(src)

            tpl = DocxTemplate(str(normalized))
            tpl.render({"nome": "JOAO TESTE"}, jinja_env=build_env())
            buf = BytesIO()
            tpl.save(buf)
            buf.seek(0)

            with ZipFile(buf) as z:
                root = ET.fromstring(z.read("word/document.xml"))

            found = False
            for run in root.findall(".//w:r", NS):
                text = "".join(node.text or "" for node in run.findall(".//w:t", NS))
                if text != "JOAO TESTE":
                    continue
                found = True
                self.assertIsNotNone(run.find("./w:rPr/w:b", NS))
                break

            self.assertTrue(found, "Texto renderizado não encontrado no documento final.")
